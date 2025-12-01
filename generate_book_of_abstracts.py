"""
IMRC 2025 Book of Abstracts Generator
=====================================
This script generates a professional Book of Abstracts document from an Excel file
containing conference submissions.

Usage:
    python generate_book_of_abstracts.py input.xlsx output.docx [track_name]

Requirements:
    pip install pandas python-docx openpyxl
"""

import sys
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============== IMRC 2025 Theme Colors ==============
IMRC_NAVY = '1A237E'
IMRC_NAVY_RGB = RGBColor(26, 35, 126)
IMRC_GRAY = RGBColor(85, 85, 85)
IMRC_BLACK = RGBColor(0, 0, 0)
IMRC_WHITE = RGBColor(255, 255, 255)


# ============== Helper Functions ==============

def clean_text(text):
    """Clean text by removing extra whitespace and normalizing line breaks."""
    if pd.isna(text):
        return ""
    text = str(text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def set_cell_border(cell, **kwargs):
    """Set border properties for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            d = kwargs[edge]
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), d.get('val', 'single'))
            el.set(qn('w:sz'), str(d.get('sz', 4)))
            el.set(qn('w:color'), d.get('color', '000000'))
            el.set(qn('w:space'), '0')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """Set background shading color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set internal margins for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for n, v in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{n}')
        m.set(qn('w:w'), str(v))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def remove_paragraph_spacing(para):
    """Remove space before and after paragraph."""
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_bookmark(para, name, bid):
    """Add a bookmark to a paragraph for internal document navigation."""
    run = para.runs[0] if para.runs else para.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bid))
    start.set(qn('w:name'), name)
    tag.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bid))
    tag.append(end)


def add_hyperlink(para, text, anchor, font='Calibri', size=11, bold=False, color='1A237E'):
    """Add a hyperlink to an internal bookmark."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), font)
    rF.set(qn('w:hAnsi'), font)
    rPr.append(rF)
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    
    if bold:
        rPr.append(OxmlElement('w:b'))
    
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)


# ============== Main Document Generator ==============

def create_book_of_abstracts(df_oral, track_title="FAE"):
    """
    Create a complete Book of Abstracts document.
    
    Args:
        df_oral: DataFrame containing oral presentations with columns:
                 'Submission ID', 'Title', 'Authors', 'Abstract'
        track_title: Name of the conference track
    
    Returns:
        Document object ready to be saved
    """
    doc = Document()
    bid = [0]
    
    def next_bid():
        bid[0] += 1
        return bid[0]
    
    # Page setup
    for sec in doc.sections:
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
    
    # ========== TITLE PAGE ==========
    for _ in range(4):
        doc.add_paragraph()
    
    # IMRC 2025
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IMRC 2025")
    r.bold = True
    r.font.name = 'Calibri Light'
    r.font.size = Pt(56)
    r.font.color.rgb = IMRC_NAVY_RGB
    
    doc.add_paragraph()
    
    # India Management Research Conference
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("India Management Research Conference")
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.color.rgb = IMRC_NAVY_RGB
    
    doc.add_paragraph()
    
    # IIM Ahmedabad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IIM Ahmedabad")
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.color.rgb = IMRC_GRAY
    
    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("December 5-7, 2025")
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.color.rgb = IMRC_GRAY
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Book of Abstracts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Book of Abstracts")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(36)
    r.font.color.rgb = IMRC_BLACK
    
    doc.add_paragraph()
    
    # Track
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Track: {track_title}")
    r.font.name = 'Calibri'
    r.font.size = Pt(20)
    r.font.color.rgb = IMRC_NAVY_RGB
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table of Contents")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(24)
    r.font.color.rgb = IMRC_NAVY_RGB
    add_bookmark(p, "TOC", next_bid())
    
    doc.add_paragraph()
    
    # TOC Table
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    hdrs = ["Submission\nID", "Title", "Authors"]
    widths = [Inches(0.8), Inches(3.5), Inches(2.5)]
    hcells = tbl.rows[0].cells
    
    for cell, hdr, w in zip(hcells, hdrs, widths):
        cell.width = w
        para = cell.paragraphs[0]
        # Handle line break in header
        if "\n" in hdr:
            parts = hdr.split("\n")
            for i, part in enumerate(parts):
                run = para.add_run(part)
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = IMRC_WHITE
                if i < len(parts) - 1:
                    para.add_run("\n")
        else:
            run = para.add_run(hdr)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = IMRC_WHITE
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, IMRC_NAVY)
        set_cell_margins(cell, 40, 40, 40, 40)
        remove_paragraph_spacing(para)
        b = {'val': 'single', 'sz': 6, 'color': IMRC_NAVY}
        set_cell_border(cell, top=b, bottom=b, left=b, right=b)
    
    # TOC rows
    for idx, (_, row) in enumerate(df_oral.iterrows()):
        sub_id = row['Submission ID']
        title = clean_text(row['Title'])
        authors = clean_text(row['Authors'])
        dr = tbl.add_row()
        cells = dr.cells
        
        # Submission ID (with bookmark for back-navigation)
        cells[0].width = Inches(0.8)
        cells[0].paragraphs[0].clear()
        add_hyperlink(cells[0].paragraphs[0], str(sub_id), f"SUB_{sub_id}", 'Calibri', 10)
        add_bookmark(cells[0].paragraphs[0], f"TOC_SUB_{sub_id}", next_bid())
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title (full title, no truncation)
        cells[1].width = Inches(3.5)
        cells[1].paragraphs[0].clear()
        add_hyperlink(cells[1].paragraphs[0], title, f"SUB_{sub_id}", 'Calibri', 10)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Authors (full names, no truncation)
        cells[2].width = Inches(2.5)
        cells[2].paragraphs[0].clear()
        add_hyperlink(cells[2].paragraphs[0], authors, f"SUB_{sub_id}", 'Calibri', 10)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Alternating row colors
        rc = 'F0F4FF' if idx % 2 == 0 else 'FFFFFF'
        b = {'val': 'single', 'sz': 4, 'color': 'D0D0D0'}
        for cell in cells:
            set_cell_shading(cell, rc)
            set_cell_margins(cell, 50, 50, 60, 60)
            set_cell_border(cell, top=b, bottom=b, left=b, right=b)
    
    doc.add_page_break()
    
    # ========== ABSTRACTS ==========
    for idx, (_, row) in enumerate(df_oral.iterrows()):
        sub_id = row['Submission ID']
        title = clean_text(row['Title'])
        authors = clean_text(row['Authors'])
        abstract = clean_text(row['Abstract'])
        
        # Abstract table
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Header row
        c1 = tbl.cell(0, 0)
        c2 = tbl.cell(0, 1)
        c1.width = Inches(3.4)
        c2.width = Inches(3.4)
        
        # Submission ID
        c1.paragraphs[0].clear()
        r = c1.paragraphs[0].add_run(f"Submission ID: {sub_id}")
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = IMRC_WHITE
        add_bookmark(c1.paragraphs[0], f"SUB_{sub_id}", next_bid())
        
        # Track
        c2.paragraphs[0].clear()
        r = c2.paragraphs[0].add_run(f"Track: {track_title}")
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = IMRC_WHITE
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        set_cell_shading(c1, IMRC_NAVY)
        set_cell_shading(c2, IMRC_NAVY)
        set_cell_margins(c1, 40, 40, 100, 100)
        set_cell_margins(c2, 40, 40, 100, 100)
        remove_paragraph_spacing(c1.paragraphs[0])
        remove_paragraph_spacing(c2.paragraphs[0])
        
        bn = {'val': 'single', 'sz': 10, 'color': IMRC_NAVY}
        set_cell_border(c1, top=bn, left=bn, bottom=bn)
        set_cell_border(c2, top=bn, right=bn, bottom=bn)
        
        # Title row
        tr = tbl.add_row()
        tc = tr.cells[0]
        tc.merge(tr.cells[1])
        tc.paragraphs[0].clear()
        r = tc.paragraphs[0].add_run(title)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.color.rgb = IMRC_NAVY_RGB
        tc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        remove_paragraph_spacing(tc.paragraphs[0])
        set_cell_border(tc, left=bn, right=bn)
        
        # Authors row
        ar = tbl.add_row()
        ac = ar.cells[0]
        ac.merge(ar.cells[1])
        ac.paragraphs[0].clear()
        r = ac.paragraphs[0].add_run(authors)
        r.italic = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = IMRC_GRAY
        ac.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(ac, 'F5F7FC')
        remove_paragraph_spacing(ac.paragraphs[0])
        set_cell_border(ac, left=bn, right=bn)
        
        # Abstract row
        abr = tbl.add_row()
        abc = abr.cells[0]
        abc.merge(abr.cells[1])
        abp = abc.paragraphs[0]
        abp.clear()
        
        r = abp.add_run("Abstract: ")
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = IMRC_NAVY_RGB
        
        r = abp.add_run(abstract)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = IMRC_BLACK
        abp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        set_cell_margins(abc, 80, 80, 100, 100)
        set_cell_border(abc, left=bn, right=bn, bottom=bn)
        
        # Back to TOC (links to specific row in TOC)
        np = doc.add_paragraph()
        np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_hyperlink(np, "↑ Back to Contents", f"TOC_SUB_{sub_id}", 'Calibri', 9, color=IMRC_NAVY)
        
        # Spacing between abstracts
        if idx < len(df_oral) - 1:
            sp = doc.add_paragraph()
            sp.add_run().font.size = Pt(8)
    
    return doc


def main():
    """Main entry point for the script."""
    # Default values
    input_file = "IMRC2025_submissions.xlsx"
    output_file = "IMRC2025_Book_of_Abstracts.docx"
    track_title = "FAE"
    
    # Parse command line arguments
    if len(sys.argv) >= 2:
        input_file = sys.argv[1]
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    if len(sys.argv) >= 4:
        track_title = sys.argv[3]
    
    print("=" * 60)
    print("IMRC 2025 Book of Abstracts Generator")
    print("=" * 60)
    print(f"\nInput file:  {input_file}")
    print(f"Output file: {output_file}")
    print(f"Track:       {track_title}")
    print()
    
    # Read Excel file
    print("Reading Excel file...")
    try:
        df = pd.read_excel(input_file)
    except FileNotFoundError:
        print(f"ERROR: File not found: {input_file}")
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Could not read Excel file: {e}")
        sys.exit(1)
    
    # Check required columns
    required_cols = ['Submission ID', 'Title', 'Authors', 'Abstract', 'Decision']
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        print(f"ERROR: Missing required columns: {missing_cols}")
        print(f"Available columns: {df.columns.tolist()}")
        sys.exit(1)
    
    # Filter for Oral Presentations
    df_oral = df[df['Decision'] == 'Oral Presentation'].copy()
    print(f"Found {len(df_oral)} Oral Presentations")
    
    if len(df_oral) == 0:
        print("ERROR: No oral presentations found in the data.")
        sys.exit(1)
    
    # Generate document
    print("\nGenerating Book of Abstracts...")
    doc = create_book_of_abstracts(df_oral, track_title)
    
    # Save document
    print(f"Saving to {output_file}...")
    try:
        doc.save(output_file)
    except PermissionError:
        print(f"ERROR: Cannot save file. Please close {output_file} if it's open.")
        sys.exit(1)
    
    print("\n" + "=" * 60)
    print("✅ SUCCESS!")
    print("=" * 60)
    print(f"\nGenerated: {output_file}")
    print(f"Total abstracts: {len(df_oral)}")
    print("\nDocument includes:")
    print("  • Title page with IMRC 2025 branding")
    print("  • Table of Contents with hyperlinks")
    print("  • All abstracts with navigation links")


if __name__ == "__main__":
    main()
